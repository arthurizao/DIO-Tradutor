import os
import requests
from docx import Document
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI
from dotenv import load_dotenv

load_dotenv()

# Variáveis de configuração
chave_api = os.getenv("YOUR_SUBSCRIPTION_KEY")
url_servico = "https://api.cognitive.microsofttranslator.com"
regiao = "eastus2"
idioma_destino = "pt-br"

def traduzir_documento(caminho_arquivo):
    doc = Document(caminho_arquivo)
    conteudo_traduzido = [traduzir_texto(paragrafo.text, idioma_destino) for paragrafo in doc.paragraphs]

    novo_doc = Document()
    for linha in conteudo_traduzido:
        novo_doc.add_paragraph(linha)

    arquivo_traduzido = caminho_arquivo.replace(".docx", f"_{idioma_destino}.docx")
    novo_doc.save(arquivo_traduzido)
    return arquivo_traduzido

def traduzir_texto(texto, para_idioma):
    caminho = '/translate'
    url = url_servico + caminho
    cabecalhos = {
        'Ocp-Apim-Subscription-Key': chave_api,
        'Ocp-Apim-Subscription-Region': regiao,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    corpo = [{'text': texto}]
    parametros = {
        'api-version': '3.0',
        'from': 'en',
        'to': para_idioma
    }

    try:
        resposta = requests.post(url, params=parametros, headers=cabecalhos, json=corpo)
        resultado = resposta.json()
        return resultado[0]["translations"][0]["text"]
    except Exception as erro:
        print(f"Erro na tradução: {erro}")
        return None

def extrair_e_limpar_texto(url):
    resposta = requests.get(url)
    if resposta.status_code == 200:
        pagina = BeautifulSoup(resposta.text, 'html.parser')
        for tag in pagina(['script', 'style']):
            tag.decompose()
        texto_bruto = pagina.get_text(separator=' ')
        linhas = (linha.strip() for linha in texto_bruto.splitlines())
        palavras = (frase.strip() for linha in linhas for frase in linha.split(" "))
        texto_limpo = '\n'.join(palavra for palavra in palavras if palavra)
        return texto_limpo
    else:
        print(f"Erro ao buscar a URL: {resposta.status_code}")
        return None

def traduzir_conteudo(conteudo, idioma):
    cliente_chat = AzureChatOpenAI(
        azure_endpoint="seu-endpoint-azure",
        api_key=os.getenv("AZURE_API_KEY"),
        api_version="2024-02-15-preview",
        deployment_name="gpt-4o-mini",
        max_retries=0
    )

    mensagem = [
        ("system", "Você é um tradutor de textos."),
        ("user", f"Traduza o seguinte texto para {idioma}: {conteudo}")
    ]

    try:
        resposta = cliente_chat.chat(mensagem)
        return resposta['choices'][0]['message']['content']
    except Exception as e:
        print(f"Erro com Azure OpenAI: {e}")
        return None

if __name__ == '__main__':
    url = 'https://www.nytimes.com/wirecutter/reviews/hanna-andersson-pajamas-review/'
    texto_extraido = extrair_e_limpar_texto(url)

    if texto_extraido:
        artigo_traduzido = traduzir_conteudo(texto_extraido, "pt-br")
        print(artigo_traduzido)

    caminho_documento = "arquivo-teste.docx"
    documento_traduzido = traduzir_documento(caminho_documento)
    print(f"Arquivo traduzido salvo em: {documento_traduzido}")

    texto_traduzido = traduzir_texto("I know you're somewhere out there, somewhere far away", idioma_destino)
    print(texto_traduzido)
